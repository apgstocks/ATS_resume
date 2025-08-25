import fitz  # PyMuPDF
from docx import Document
import re
import os
from typing import Dict, List, Optional
import logging

logger = logging.getLogger(__name__)

class ResumeParser:
    def __init__(self):
        # Common resume section patterns
        self.section_patterns = {
            'contact': [
                r'contact\s*information',
                r'personal\s*information',
                r'contact\s*details'
            ],
            'summary': [
                r'professional\s*summary',
                r'career\s*summary',
                r'summary',
                r'objective',
                r'profile'
            ],
            'experience': [
                r'work\s*experience',
                r'professional\s*experience',
                r'employment\s*history',
                r'experience',
                r'career\s*history'
            ],
            'education': [
                r'education',
                r'academic\s*background',
                r'qualifications'
            ],
            'skills': [
                r'skills',
                r'technical\s*skills',
                r'core\s*competencies',
                r'areas\s*of\s*expertise'
            ],
            'certifications': [
                r'certifications',
                r'certificates',
                r'licenses',
                r'professional\s*certifications'
            ]
        }
        
        # Email and phone patterns
        self.email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        self.phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            
            # Check if we extracted meaningful text
            if not text or len(text.strip()) < 10:
                raise Exception("Could not extract text from PDF. The file may be image-based or corrupted.")
                
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            # Try to provide more helpful error message
            if "cannot open broken document" in str(e):
                raise Exception("Invalid or corrupted PDF file. Please ensure the file is a valid PDF document.")
            elif "password" in str(e).lower():
                raise Exception("Password-protected PDF files are not supported. Please upload an unprotected PDF.")
            else:
                raise Exception(f"Failed to parse PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text).strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise Exception(f"Failed to parse DOCX: {str(e)}")
    
    def extract_text_from_doc(self, file_path: str) -> str:
        """Extract text from DOC file (legacy format)"""
        try:
            # For legacy DOC files, we'll try to use python-docx
            # Note: python-docx doesn't support old .doc format well
            # In production, consider using python-docx2txt or converting with LibreOffice
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text).strip()
        except Exception as e:
            logger.warning(f"Could not parse legacy DOC file: {str(e)}")
            # Return a message indicating the format isn't fully supported
            raise Exception("Legacy .DOC format not fully supported. Please convert to PDF or DOCX.")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            raise Exception(f"Failed to parse TXT: {str(e)}")
    
    def parse_resume(self, file_path: str, file_type: str) -> Dict:
        """Main parsing function that routes to appropriate parser"""
        try:
            if file_type.lower() == 'pdf':
                text = self.extract_text_from_pdf(file_path)
            elif file_type.lower() == 'docx':
                text = self.extract_text_from_docx(file_path)
            elif file_type.lower() == 'doc':
                text = self.extract_text_from_doc(file_path)
            elif file_type.lower() == 'txt':
                text = self.extract_text_from_txt(file_path)
            else:
                raise Exception(f"Unsupported file type: {file_type}")
            
            if not text or len(text.strip()) < 50:
                raise Exception("Could not extract meaningful text from resume. Please check if the file is corrupted or text-based.")
            
            # Parse the extracted text
            parsed_data = self.analyze_resume_structure(text)
            parsed_data['raw_text'] = text
            parsed_data['file_type'] = file_type
            
            return parsed_data
            
        except Exception as e:
            logger.error(f"Error parsing resume: {str(e)}")
            raise e
        finally:
            # Clean up the uploaded file
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass
    
    def analyze_resume_structure(self, text: str) -> Dict:
        """Analyze resume structure and extract sections"""
        # Clean the text
        cleaned_text = self.clean_text(text)
        lines = cleaned_text.split('\n')
        
        # Extract contact information
        contact_info = self.extract_contact_info(cleaned_text)
        
        # Detect sections
        sections = self.detect_sections(lines)
        
        # Extract key information
        skills = self.extract_skills(cleaned_text)
        
        return {
            'contact_info': contact_info,
            'sections': sections,
            'skills': skills,
            'word_count': len(cleaned_text.split()),
            'line_count': len(lines),
            'has_bullet_points': self.has_bullet_points(cleaned_text),
            'formatting_issues': self.detect_formatting_issues(cleaned_text)
        }
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters that might interfere with parsing
        text = re.sub(r'[^\w\s@.-]', ' ', text)
        # Normalize line breaks
        text = re.sub(r'\n\s*\n', '\n', text)
        return text.strip()
    
    def extract_contact_info(self, text: str) -> Dict:
        """Extract contact information from resume"""
        contact = {
            'email': None,
            'phone': None,
            'has_contact_section': False
        }
        
        # Extract email
        email_matches = re.findall(self.email_pattern, text, re.IGNORECASE)
        if email_matches:
            contact['email'] = email_matches[0]
        
        # Extract phone
        phone_matches = re.findall(self.phone_pattern, text)
        if phone_matches:
            contact['phone'] = phone_matches[0]
        
        # Check if there's a dedicated contact section
        for pattern in self.section_patterns['contact']:
            if re.search(pattern, text, re.IGNORECASE):
                contact['has_contact_section'] = True
                break
        
        return contact
    
    def detect_sections(self, lines: List[str]) -> Dict:
        """Detect different sections in the resume"""
        sections = {}
        current_section = None
        
        for section_name, patterns in self.section_patterns.items():
            sections[section_name] = {
                'present': False,
                'content': [],
                'line_number': None
            }
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if not line_lower:
                continue
            
            # Check if this line is a section header
            for section_name, patterns in self.section_patterns.items():
                for pattern in patterns:
                    if re.match(f'^{pattern}$', line_lower) or re.match(f'^{pattern}:', line_lower):
                        sections[section_name]['present'] = True
                        sections[section_name]['line_number'] = i
                        current_section = section_name
                        break
                if current_section == section_name:
                    break
            
            # Add content to current section
            if current_section and line.strip():
                sections[current_section]['content'].append(line.strip())
        
        return sections
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text"""
        # Common technical skills patterns
        tech_skills = [
            'python', 'java', 'javascript', 'react', 'node.js', 'angular', 'vue.js',
            'html', 'css', 'sql', 'mongodb', 'postgresql', 'mysql', 'aws', 'azure',
            'docker', 'kubernetes', 'git', 'jenkins', 'ci/cd', 'agile', 'scrum',
            'machine learning', 'data analysis', 'pandas', 'numpy', 'tensorflow',
            'pytorch', 'scikit-learn', 'tableau', 'power bi', 'excel', 'powerpoint'
        ]
        
        found_skills = []
        text_lower = text.lower()
        
        for skill in tech_skills:
            if skill in text_lower:
                found_skills.append(skill.title())
        
        return found_skills
    
    def has_bullet_points(self, text: str) -> bool:
        """Check if resume uses bullet points"""
        bullet_patterns = [r'•', r'●', r'▪', r'◦', r'-\s', r'\*\s']
        for pattern in bullet_patterns:
            if re.search(pattern, text):
                return True
        return False
    
    def detect_formatting_issues(self, text: str) -> List[str]:
        """Detect potential formatting issues that might affect ATS parsing"""
        issues = []
        
        # Check for tables (common ATS issue)
        if '|' in text or '\t' in text:
            issues.append("Potential table formatting detected")
        
        # Check for excessive special characters
        special_char_count = len(re.findall(r'[^\w\s@.-]', text))
        if special_char_count > len(text.split()) * 0.1:  # More than 10% special chars
            issues.append("Excessive special characters detected")
        
        # Check for very long lines (might indicate formatting issues)
        lines = text.split('\n')
        long_lines = [line for line in lines if len(line) > 100]
        if len(long_lines) > len(lines) * 0.3:  # More than 30% long lines
            issues.append("Potential formatting issues with line breaks")
        
        return issues