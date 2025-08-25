import fitz  # PyMuPDF
from docx import Document
import re
import os
import tempfile
from typing import Dict, List, Tuple, Optional
import nltk
import spacy
import textstat
from collections import Counter
import logging

logger = logging.getLogger(__name__)

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
    nltk.data.find('taggers/averaged_perceptron_tagger')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('averaged_perceptron_tagger')

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    logger.warning("spaCy model not found, some features may be limited")
    nlp = None

class AdvancedResumeParser:
    def __init__(self):
        self.ats_friendly_fonts = [
            'arial', 'calibri', 'times new roman', 'helvetica', 
            'georgia', 'trebuchet ms', 'verdana'
        ]
        
        self.standard_sections = [
            'work experience', 'professional experience', 'experience', 'employment',
            'education', 'academic background', 'qualifications',
            'skills', 'technical skills', 'core competencies', 'technologies',
            'certifications', 'certificates', 'licenses',
            'summary', 'professional summary', 'objective', 'profile'
        ]
        
        self.action_verbs = [
            'achieved', 'accomplished', 'advanced', 'analyzed', 'built', 'created',
            'delivered', 'developed', 'enhanced', 'established', 'executed',
            'generated', 'improved', 'increased', 'initiated', 'launched',
            'led', 'managed', 'optimized', 'organized', 'performed', 'planned',
            'produced', 'reduced', 'resolved', 'streamlined', 'supervised',
            'transformed', 'utilized', 'collaborated', 'coordinated', 'implemented'
        ]
        
        self.tech_skills = [
            # Programming Languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby',
            'go', 'rust', 'scala', 'kotlin', 'swift', 'r', 'matlab', 'sql',
            
            # Web Technologies
            'html', 'css', 'react', 'angular', 'vue.js', 'node.js', 'express.js',
            'django', 'flask', 'laravel', 'spring boot', 'asp.net',
            
            # Databases
            'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'oracle',
            'sqlite', 'cassandra', 'dynamodb',
            
            # Cloud & DevOps
            'aws', 'azure', 'google cloud', 'docker', 'kubernetes', 'jenkins',
            'gitlab ci', 'github actions', 'terraform', 'ansible',
            
            # Tools & Technologies
            'git', 'jira', 'confluence', 'tableau', 'power bi', 'excel',
            'photoshop', 'figma', 'sketch', 'salesforce', 'hubspot'
        ]
        
        self.soft_skills = [
            'leadership', 'management', 'communication', 'teamwork', 'collaboration',
            'problem solving', 'analytical', 'creative', 'organized', 'adaptable',
            'detail-oriented', 'time management', 'critical thinking', 'negotiation',
            'presentation', 'customer service', 'project management'
        ]
        
        self.certifications = [
            'pmp', 'aws certified', 'azure certified', 'google certified',
            'cissp', 'cisa', 'cism', 'comptia', 'cisco certified', 'microsoft certified',
            'salesforce certified', 'scrum master', 'six sigma', 'itil'
        ]

    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from PDF and analyze formatting"""
        try:
            doc = fitz.open(file_path)
            text = ""
            formatting_info = {
                'total_pages': len(doc),
                'has_images': False,
                'has_tables': False,
                'fonts_used': set(),
                'formatting_issues': []
            }
            
            for page in doc:
                # Extract text
                page_text = page.get_text()
                text += page_text + "\n"
                
                # Check for images
                if page.get_images():
                    formatting_info['has_images'] = True
                
                # Check for tables (simple heuristic)
                if '\t' in page_text or '|' in page_text:
                    formatting_info['has_tables'] = True
                
                # Extract font information
                blocks = page.get_text("dict")["blocks"]
                for block in blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                font_name = span.get("font", "").lower()
                                if font_name:
                                    formatting_info['fonts_used'].add(font_name)
            
            doc.close()
            
            # Check for ATS-friendly fonts
            if formatting_info['fonts_used']:
                non_ats_fonts = [f for f in formatting_info['fonts_used'] 
                               if not any(ats_font in f for ats_font in self.ats_friendly_fonts)]
                if non_ats_fonts:
                    formatting_info['formatting_issues'].append(
                        f"Non-ATS friendly fonts detected: {', '.join(list(non_ats_fonts)[:3])}"
                    )
            
            if formatting_info['has_images']:
                formatting_info['formatting_issues'].append("Contains images which may not be ATS-friendly")
            
            if formatting_info['has_tables']:
                formatting_info['formatting_issues'].append("Contains tables which may cause parsing issues")
            
            return text.strip(), formatting_info
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise Exception(f"Failed to parse PDF: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from DOCX and analyze formatting"""
        try:
            doc = Document(file_path)
            text = []
            formatting_info = {
                'has_images': False,
                'has_tables': len(doc.tables) > 0,
                'fonts_used': set(),
                'formatting_issues': []
            }
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
                
                # Check fonts
                for run in paragraph.runs:
                    if run.font.name:
                        formatting_info['fonts_used'].add(run.font.name.lower())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            
            # Check for images
            from docx.oxml.ns import nsdecls, qn
            from docx.oxml import parse_xml
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    formatting_info['has_images'] = True
                    break
            
            # Check formatting issues
            if formatting_info['has_tables']:
                formatting_info['formatting_issues'].append("Contains tables which may cause ATS parsing issues")
            
            if formatting_info['has_images']:
                formatting_info['formatting_issues'].append("Contains images which are not ATS-friendly")
            
            # Check fonts
            if formatting_info['fonts_used']:
                non_ats_fonts = [f for f in formatting_info['fonts_used'] 
                               if not any(ats_font in f for ats_font in self.ats_friendly_fonts)]
                if non_ats_fonts:
                    formatting_info['formatting_issues'].append(
                        f"Non-ATS friendly fonts: {', '.join(list(non_ats_fonts)[:3])}"
                    )
            
            return '\n'.join(text).strip(), formatting_info
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise Exception(f"Failed to parse DOCX: {str(e)}")

    def parse_resume(self, file_path: str, file_type: str) -> Dict:
        """Main parsing function"""
        try:
            if file_type.lower() == 'pdf':
                text, formatting_info = self.extract_text_from_pdf(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                text, formatting_info = self.extract_text_from_docx(file_path)
            else:
                # For TXT files
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read().strip()
                formatting_info = {'formatting_issues': [], 'fonts_used': set(), 'has_images': False, 'has_tables': False}
            
            if not text or len(text.strip()) < 100:
                raise Exception("Could not extract meaningful text from resume. Please ensure the file contains readable text.")
            
            # Comprehensive analysis
            parsed_data = {
                'raw_text': text,
                'file_type': file_type,
                'formatting_info': formatting_info,
                'contact_info': self.extract_contact_info(text),
                'sections': self.detect_sections(text),
                'work_experience': self.analyze_work_experience(text),
                'education': self.extract_education(text),
                'skills': self.extract_skills(text),
                'certifications': self.extract_certifications(text),
                'readability': self.analyze_readability(text),
                'grammar_issues': self.check_grammar_basic(text),
                'ats_parsing_test': self.simulate_ats_parsing(text)
            }
            
            return parsed_data
            
        except Exception as e:
            logger.error(f"Error parsing resume: {str(e)}")
            raise e
        finally:
            # Clean up uploaded file
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass

    def extract_contact_info(self, text: str) -> Dict:
        """Extract contact information"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'(\+?1?[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        
        emails = re.findall(email_pattern, text)
        phones = re.findall(phone_pattern, text)
        linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
        
        # Check for bias-sensitive information
        bias_keywords = ['age', 'married', 'single', 'gender', 'race', 'religion', 'photo']
        bias_issues = [keyword for keyword in bias_keywords if keyword.lower() in text.lower()]
        
        return {
            'emails': emails,
            'phones': ['-'.join(phone[1:]) for phone in phones],
            'linkedin': linkedin,
            'has_email': len(emails) > 0,
            'has_phone': len(phones) > 0,
            'has_linkedin': len(linkedin) > 0,
            'bias_issues': bias_issues
        }

    def detect_sections(self, text: str) -> Dict:
        """Detect resume sections"""
        lines = text.split('\n')
        sections_found = {}
        
        # Initialize all sections as not found
        section_map = {
            'work_experience': ['work experience', 'professional experience', 'experience', 'employment', 'career'],
            'education': ['education', 'academic background', 'qualifications', 'degree'],
            'skills': ['skills', 'technical skills', 'competencies', 'technologies', 'expertise'],
            'certifications': ['certifications', 'certificates', 'licenses', 'credentials'],
            'summary': ['summary', 'professional summary', 'objective', 'profile', 'about'],
            'contact': ['contact', 'personal information']
        }
        
        for section_key, keywords in section_map.items():
            sections_found[section_key] = {
                'found': False,
                'header_line': None,
                'content_lines': []
            }
            
            for i, line in enumerate(lines):
                line_lower = line.strip().lower()
                if any(keyword in line_lower for keyword in keywords):
                    sections_found[section_key]['found'] = True
                    sections_found[section_key]['header_line'] = i
                    
                    # Extract content (next 10 lines or until next section)
                    content_lines = []
                    for j in range(i + 1, min(i + 11, len(lines))):
                        if lines[j].strip():
                            content_lines.append(lines[j].strip())
                    sections_found[section_key]['content_lines'] = content_lines
                    break
        
        return sections_found

    def analyze_work_experience(self, text: str) -> Dict:
        """Analyze work experience section"""
        text_lower = text.lower()
        
        # Check for quantified achievements
        numbers_pattern = r'\b\d+(?:\.\d+)?(?:%|k|million|billion|\$|,\d{3})*\b'
        quantified_achievements = len(re.findall(numbers_pattern, text))
        
        # Count action verbs
        action_verb_count = sum(1 for verb in self.action_verbs if verb in text_lower)
        
        # Check for reverse chronological order (look for years)
        years = re.findall(r'\b(19|20)\d{2}\b', text)
        is_chronological = len(years) >= 2 and years == sorted(years, reverse=True)
        
        # Check for job titles
        common_titles = [
            'manager', 'director', 'engineer', 'developer', 'analyst', 'specialist',
            'coordinator', 'assistant', 'supervisor', 'lead', 'senior', 'junior'
        ]
        job_titles_found = [title for title in common_titles if title in text_lower]
        
        # Check for company names (look for "at", "with", "|" patterns)
        company_indicators = len(re.findall(r'\s+at\s+\w+|\s+with\s+\w+|\w+\s*\|', text, re.IGNORECASE))
        
        return {
            'quantified_achievements': quantified_achievements,
            'action_verbs_count': action_verb_count,
            'is_reverse_chronological': is_chronological,
            'job_titles_found': job_titles_found,
            'company_indicators': company_indicators,
            'years_mentioned': len(years)
        }

    def extract_education(self, text: str) -> Dict:
        """Extract education information"""
        text_lower = text.lower()
        
        # Degree patterns
        degree_patterns = [
            r'\b(bachelor|master|phd|doctorate|associate|diploma)\b',
            r'\b(b\.?[sa]\.?|m\.?[sa]\.?|ph\.?d\.?|m\.?b\.?a\.?)\b',
            r'\b(undergraduate|graduate|postgraduate)\b'
        ]
        
        degrees = []
        for pattern in degree_patterns:
            degrees.extend(re.findall(pattern, text_lower))
        
        # University/College patterns
        institution_keywords = ['university', 'college', 'institute', 'school', 'academy']
        institutions = []
        for keyword in institution_keywords:
            if keyword in text_lower:
                institutions.append(keyword)
        
        # GPA pattern
        gpa_pattern = r'gpa\s*:?\s*([0-9]\.[0-9])'
        gpa_matches = re.findall(gpa_pattern, text_lower)
        
        return {
            'degrees_found': list(set(degrees)),
            'institutions_mentioned': len(institutions) > 0,
            'gpa_mentioned': len(gpa_matches) > 0,
            'gpa_values': gpa_matches
        }

    def extract_skills(self, text: str) -> Dict:
        """Extract and categorize skills"""
        text_lower = text.lower()
        
        tech_skills_found = [skill for skill in self.tech_skills if skill in text_lower]
        soft_skills_found = [skill for skill in self.soft_skills if skill in text_lower]
        
        # Calculate skill diversity
        total_skills = len(tech_skills_found) + len(soft_skills_found)
        
        return {
            'technical_skills': tech_skills_found,
            'soft_skills': soft_skills_found,
            'total_skills_count': total_skills,
            'tech_skills_count': len(tech_skills_found),
            'soft_skills_count': len(soft_skills_found),
            'skills_diversity_score': min(100, total_skills * 5)
        }

    def extract_certifications(self, text: str) -> Dict:
        """Extract certifications"""
        text_lower = text.lower()
        
        certifications_found = []
        for cert in self.certifications:
            if cert in text_lower:
                certifications_found.append(cert)
        
        return {
            'certifications_found': certifications_found,
            'certifications_count': len(certifications_found)
        }

    def analyze_readability(self, text: str) -> Dict:
        """Analyze text readability"""
        try:
            flesch_reading_ease = textstat.flesch_reading_ease(text)
            flesch_grade = textstat.flesch_kincaid_grade(text)
            word_count = len(text.split())
            sentence_count = len(re.findall(r'[.!?]+', text))
            
            # Calculate average sentence length
            avg_sentence_length = word_count / max(sentence_count, 1)
            
            return {
                'flesch_reading_ease': flesch_reading_ease,
                'flesch_grade_level': flesch_grade,
                'word_count': word_count,
                'sentence_count': sentence_count,
                'avg_sentence_length': avg_sentence_length,
                'readability_score': min(100, max(0, flesch_reading_ease))
            }
        except:
            return {
                'flesch_reading_ease': 50,
                'flesch_grade_level': 8,
                'word_count': len(text.split()),
                'sentence_count': len(re.findall(r'[.!?]+', text)),
                'avg_sentence_length': 15,
                'readability_score': 50
            }

    def check_grammar_basic(self, text: str) -> List[str]:
        """Basic grammar and spelling check"""
        issues = []
        
        # Check for common issues
        if re.search(r'\s{2,}', text):
            issues.append("Multiple consecutive spaces found")
        
        if re.search(r'[a-z]\.[A-Z]', text):
            issues.append("Missing space after period")
        
        # Check for incomplete sentences
        sentences = re.split(r'[.!?]+', text)
        short_sentences = [s for s in sentences if len(s.strip().split()) < 3 and len(s.strip()) > 0]
        if len(short_sentences) > len(sentences) * 0.3:
            issues.append("Many incomplete or very short sentences")
        
        return issues

    def simulate_ats_parsing(self, text: str) -> Dict:
        """Simulate how an ATS might parse the resume"""
        # Remove special characters and formatting
        cleaned_text = re.sub(r'[^\w\s@.-]', ' ', text)
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        
        # Check what might be lost in parsing
        parsing_issues = []
        
        if len(cleaned_text) < len(text) * 0.8:
            parsing_issues.append("Significant text may be lost due to special formatting")
        
        # Check for proper section parsing
        section_headers = ['experience', 'education', 'skills', 'summary']
        parseable_sections = sum(1 for header in section_headers if header in cleaned_text.lower())
        
        return {
            'original_length': len(text),
            'parsed_length': len(cleaned_text),
            'parsing_efficiency': (len(cleaned_text) / len(text)) * 100,
            'parseable_sections': parseable_sections,
            'parsing_issues': parsing_issues
        }